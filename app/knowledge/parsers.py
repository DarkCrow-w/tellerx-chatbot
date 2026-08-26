"""Format-aware document parsing into layout-preserving logical units."""

from __future__ import annotations

import csv
import html
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import ClassVar

from bs4 import BeautifulSoup
from docx import Document as WordDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from pypdf import PdfReader

from app.knowledge.chunking import ParsedUnit


class UnsupportedDocument(ValueError):
    """文件格式不在支持范围内，或缺少必要的转换工具。"""


class EmptyDocument(ValueError):
    """文件中没有可提取的文字内容。"""


class DocumentParser:
    """按文件格式解析文档，同时尽可能保留标题、页码和表格位置。"""

    revision: ClassVar[str] = "v3"
    allowed_suffixes: ClassVar[set[str]] = {
        ".docx",
        ".doc",
        ".xlsx",
        ".xls",
        ".xlsm",
        ".md",
        ".markdown",
        ".txt",
        ".html",
        ".htm",
        ".pdf",
        ".csv",
    }

    def __init__(self, backend: str = "native"):
        """记录解析后端名称，用于构造可追踪的解析器指纹。"""

        self.backend = backend

    def parse(self, path: Path) -> tuple[list[ParsedUnit], list[str]]:
        """分派到对应格式解析器，返回逻辑单元和非致命警告。"""

        suffix = path.suffix.lower()
        if suffix not in self.allowed_suffixes:
            if suffix == ".one":
                raise UnsupportedDocument(
                    "Native OneNote .one files must be exported to DOCX, PDF, or HTML"
                )
            raise UnsupportedDocument(f"Unsupported file type: {suffix}")
        if suffix in {".doc", ".xls"}:
            return self._convert_legacy(path)
        if suffix in {".xlsx", ".xlsm"}:
            return self._parse_excel(path)
        if suffix == ".csv":
            return self._parse_csv(path)

        if suffix == ".docx":
            units = self._parse_docx_native(path)
        elif suffix == ".pdf":
            units = self._parse_pdf_native(path)
        elif suffix in {".html", ".htm"}:
            units = self._parse_html(path)
        elif suffix in {".md", ".markdown"}:
            units = self._parse_markdown(path.read_text(encoding="utf-8", errors="replace"))
        else:
            units = [ParsedUnit(text=path.read_text(encoding="utf-8", errors="replace"))]
        return self._ensure_content(units), []

    @staticmethod
    def _ensure_content(units: list[ParsedUnit]) -> list[ParsedUnit]:
        """过滤空单元，并对扫描件等无文本文件给出明确错误。"""

        cleaned = [unit for unit in units if unit.text.strip()]
        if not cleaned:
            raise EmptyDocument("No extractable text was found; scanned PDFs require OCR")
        return cleaned

    @staticmethod
    def _parse_markdown(text: str) -> list[ParsedUnit]:
        """按标题层级切分 Markdown，并把完整标题路径附到正文。"""

        units: list[ParsedUnit] = []
        headings: list[str] = []
        buffer: list[str] = []

        def flush() -> None:
            """把当前标题下积累的正文提交为一个逻辑单元。"""

            if buffer:
                units.append(ParsedUnit(text="\n".join(buffer), heading_path=" > ".join(headings)))
                buffer.clear()

        for line in text.splitlines():
            stripped = line.strip()
            if stripped.startswith("#") and stripped.lstrip("#").startswith(" "):
                flush()
                level = len(stripped) - len(stripped.lstrip("#"))
                title = stripped[level:].strip()
                headings[:] = headings[: level - 1]
                headings.append(title)
            else:
                buffer.append(line)
        flush()
        return units

    @staticmethod
    def _parse_docx_native(path: Path) -> list[ParsedUnit]:
        """按 Word 原始块顺序解析段落和表格，保留标题层级。"""

        document = WordDocument(path)
        units: list[ParsedUnit] = []
        headings: list[str] = []
        paragraphs: list[str] = []
        table_index = 0

        def flush() -> None:
            """在标题或表格边界处提交已积累的 Word 段落。"""

            if paragraphs:
                units.append(
                    ParsedUnit(text="\n".join(paragraphs), heading_path=" > ".join(headings))
                )
                paragraphs.clear()

        for block in document.iter_inner_content():
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if not text:
                    continue
                style = block.style.name if block.style else ""
                if style.lower().startswith("heading"):
                    flush()
                    try:
                        level = int(style.split()[-1])
                    except ValueError:
                        level = 1
                    headings[:] = headings[: level - 1]
                    headings.append(text)
                else:
                    paragraphs.append(text)
            elif isinstance(block, Table):
                flush()
                table_index += 1
                rows = [
                    [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    for row in block.rows
                ]
                if rows:
                    lines = [" | ".join(row) for row in rows]
                    units.append(
                        ParsedUnit(
                            text="\n".join(lines),
                            heading_path=(
                                f"{' > '.join(headings)} > Table {table_index}".strip(" >")
                            ),
                            is_table=True,
                        )
                    )
        flush()
        return units

    @staticmethod
    def _parse_pdf_native(path: Path) -> list[ParsedUnit]:
        """逐页提取 PDF 文本，并修复视觉换行拆开的 ASCII 标识符。"""

        reader = PdfReader(path)
        units = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            # PDF 提取器可能在视觉换行处拆开 ASCII 标识（如 TX-0929-R\nEADY）。
            text = re.sub(
                r"((?<!\w)[A-Za-z][A-Za-z0-9-]*-[A-Za-z0-9-]+)\n(?=[A-Za-z0-9-])",
                r"\1",
                text,
            )
            if text.strip():
                units.append(
                    ParsedUnit(
                        text=text, page_number=page_number, heading_path=f"Page {page_number}"
                    )
                )
        return units

    @staticmethod
    def _parse_html(path: Path) -> list[ParsedUnit]:
        """移除非内容节点，按标题、正文和表格的可见结构解析 HTML。"""

        soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="replace"), "html.parser")
        for element in soup(["script", "style", "noscript"]):
            element.decompose()
        units: list[ParsedUnit] = []
        headings: list[str] = []
        buffer: list[str] = []

        def flush() -> None:
            """提交当前 HTML 标题层级下的可见正文。"""

            text = "\n".join(buffer).strip()
            if text:
                units.append(
                    ParsedUnit(text=html.unescape(text), heading_path=" > ".join(headings))
                )
            buffer.clear()

        for element in soup.find_all(
            [
                "h1",
                "h2",
                "h3",
                "h4",
                "h5",
                "h6",
                "p",
                "li",
                "dt",
                "dd",
                "aside",
                "pre",
                "blockquote",
                "table",
            ]
        ):
            if element.name.startswith("h"):
                flush()
                level = int(element.name[1])
                headings[:] = headings[: level - 1]
                headings.append(element.get_text(" ", strip=True))
            elif element.name == "table":
                flush()
                rows = []
                for row in element.find_all("tr"):
                    rows.append(
                        " | ".join(
                            cell.get_text(" ", strip=True) for cell in row.find_all(["th", "td"])
                        )
                    )
                if rows:
                    units.append(
                        ParsedUnit(
                            text="\n".join(rows), heading_path=" > ".join(headings), is_table=True
                        )
                    )
            else:
                text = element.get_text(" ", strip=True)
                if text:
                    buffer.append(text)
        flush()
        return units

    @staticmethod
    def _parse_excel(path: Path) -> tuple[list[ParsedUnit], list[str]]:
        """同时读取公式与缓存值，按固定行窗口输出带单元格范围的表格块。"""

        formulas = load_workbook(path, read_only=True, data_only=False, keep_links=False)
        values = load_workbook(path, read_only=True, data_only=True, keep_links=False)
        units: list[ParsedUnit] = []
        warnings: list[str] = []
        for formula_sheet in formulas.worksheets:
            if formula_sheet.sheet_state != "visible":
                continue
            value_sheet = values[formula_sheet.title]
            rows: list[tuple[int, list[str]]] = []
            max_column = 0
            for row_number, (formula_row, value_row) in enumerate(
                zip(formula_sheet.iter_rows(), value_sheet.iter_rows()), start=1
            ):
                rendered: list[str] = []
                for formula_cell, value_cell in zip(formula_row, value_row):
                    raw = formula_cell.value
                    cached = value_cell.value
                    if raw is None and cached is None:
                        rendered.append("")
                    elif isinstance(raw, str) and raw.startswith("="):
                        rendered.append(
                            f"{cached if cached is not None else '[no cached value]'} ({raw})"
                        )
                        if cached is None:
                            warnings.append(
                                f"{formula_sheet.title}!{formula_cell.coordinate} has no cached formula value"
                            )
                    else:
                        rendered.append(str(cached if cached is not None else raw))
                while rendered and not rendered[-1]:
                    rendered.pop()
                if rendered and any(cell for cell in rendered):
                    max_column = max(max_column, len(rendered))
                    rows.append((row_number, rendered))
            if not rows:
                continue
            header = rows[0][1]
            for start in range(0, len(rows), 25):
                group = rows[start : start + 25]
                lines = []
                if start > 0:
                    lines.append(" | ".join(header))
                lines.extend(" | ".join(cells) for _, cells in group)
                first_row, last_row = group[0][0], group[-1][0]
                cell_range = f"A{first_row}:{get_column_letter(max_column)}{last_row}"
                units.append(
                    ParsedUnit(
                        text="\n".join(lines),
                        heading_path=f"Worksheet: {formula_sheet.title}",
                        sheet_name=formula_sheet.title,
                        cell_range=cell_range,
                        is_table=True,
                    )
                )
        formulas.close()
        values.close()
        return DocumentParser._ensure_content(units), list(dict.fromkeys(warnings))

    @staticmethod
    def _parse_csv(path: Path) -> tuple[list[ParsedUnit], list[str]]:
        """按固定行窗口切分 CSV，并在后续窗口重复表头。"""

        with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as handle:
            rows = list(csv.reader(handle))
        units = []
        for start in range(0, len(rows), 25):
            group = rows[start : start + 25]
            if start > 0 and rows:
                group = [rows[0], *group]
            units.append(
                ParsedUnit(
                    text="\n".join(" | ".join(row) for row in group),
                    cell_range=f"row {start + 1}-{min(start + 25, len(rows))}",
                    is_table=True,
                )
            )
        return DocumentParser._ensure_content(units), []

    def _convert_legacy(self, path: Path) -> tuple[list[ParsedUnit], list[str]]:
        """通过 LibreOffice 临时转换旧版 Office 文件，再复用现代格式解析器。"""

        soffice = shutil.which("soffice")
        if not soffice:
            raise UnsupportedDocument("LibreOffice is required for legacy .doc/.xls files")
        target_format = "docx" if path.suffix.lower() == ".doc" else "xlsx"
        with tempfile.TemporaryDirectory(prefix="knowledge-convert-") as temp_dir:
            result = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    target_format,
                    "--outdir",
                    temp_dir,
                    str(path),
                ],
                capture_output=True,
                text=True,
                timeout=120,
                check=False,
            )
            converted = Path(temp_dir) / f"{path.stem}.{target_format}"
            if result.returncode != 0 or not converted.exists():
                raise UnsupportedDocument("LibreOffice could not convert the legacy Office file")
            units, warnings = self.parse(converted)
            return units, ["Legacy Office file was converted with LibreOffice", *warnings]
