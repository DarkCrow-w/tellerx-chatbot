"""Synthetic corpus generation and repeatable retrieval/answer benchmarks."""

from __future__ import annotations

import argparse
import hashlib
import json
import logging
import math
import random
import re
import shutil
import statistics
import tempfile
import time
import uuid
from collections import Counter, defaultdict
from dataclasses import asdict, dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document as WordDocument
from openpyxl import Workbook
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen.canvas import Canvas
from sqlalchemy import create_engine, delete, func, select
from sqlalchemy.orm import Session

from app.core.config import Settings
from app.core.container import answer_service, qwen_client, search_index
from app.db import Base, SessionLocal, engine
from app.db.models import (
    Chunk,
    ChunkEmbedding,
    Document,
    DocumentArtifact,
    DocumentVersion,
    EmbeddingCache,
    EmbeddingModel,
    IndexGeneration,
    IndexSyncState,
    IngestionJob,
    OutboxEvent,
    Project,
)
from app.integrations.qwen import ChatCallResult, QwenAPIError, Usage
from app.integrations.search import SearchIndex, _lexical_signals
from app.knowledge.chunking import chunk_units
from app.knowledge.parsers import DocumentParser
from app.services.answering import AnswerService
from app.services.ingestion import IngestionService
from app.services.retrieval import Retriever

logger = logging.getLogger(__name__)

KNOWLEDGE_RESET_ORDER = [
    IndexSyncState,
    IndexGeneration,
    OutboxEvent,
    ChunkEmbedding,
    EmbeddingCache,
    EmbeddingModel,
    DocumentArtifact,
    Chunk,
    IngestionJob,
    DocumentVersion,
    Document,
    Project,
]

EVIDENCE_BLOCK = re.compile(
    r"\[EVIDENCE id=(?P<id>[^\]]+)]\n"
    r"file=(?P<file>[^\n]*)\nstatus=(?P<status>[^\n]*)\n"
    r"version=(?P<version>[^\n]*)\nlocation=(?P<location>[^\n]*)\n"
    r"(?P<content>.*?)\n\[/EVIDENCE]",
    re.DOTALL,
)


class OfflineBenchmarkQwen:
    """Make scale validation deterministic and guarantee zero external API calls."""

    def embeddings(self, texts: list[str]) -> tuple[list[list[float]], Usage]:
        del texts
        raise QwenAPIError("offline benchmark", code="offline_benchmark")

    def rerank(self, query: str, documents: list[str], top_n: int) -> list[tuple[int, float]]:
        del query, documents, top_n
        raise QwenAPIError("offline benchmark", code="offline_benchmark")


def _offline_feature_vector(text: str, dimensions: int = 1024) -> list[float]:
    normalized = text.casefold()
    features = re.findall(r"[a-z0-9][a-z0-9_-]*", normalized)
    cjk_runs = re.findall(r"[\u3400-\u9fff]+", normalized)
    for run in cjk_runs:
        features.extend(run[index : index + 2] for index in range(max(1, len(run) - 1)))
    vector = [0.0] * dimensions
    for feature in features:
        digest = hashlib.blake2b(feature.encode("utf-8"), digest_size=8).digest()
        raw = int.from_bytes(digest, "big")
        vector[raw % dimensions] += 1.0 if (raw >> 10) & 1 else -1.0
    norm = math.sqrt(sum(value * value for value in vector)) or 1.0
    return [value / norm for value in vector]


class OfflineHybridQwen(OfflineBenchmarkQwen):
    """Deterministic feature hashing for mechanical k-NN/RRF/rerank tests."""

    def embeddings(self, texts: list[str]) -> tuple[list[list[float]], Usage]:
        return [_offline_feature_vector(text) for text in texts], Usage()

    def rerank(self, query: str, documents: list[str], top_n: int) -> list[tuple[int, float]]:
        query_vector = _offline_feature_vector(query)
        signals = [signal.casefold() for signal in _lexical_signals(query)]
        scored = []
        for index, document in enumerate(documents):
            vector = _offline_feature_vector(document)
            score = sum(left * right for left, right in zip(query_vector, vector))
            normalized_document = document.casefold()
            score += 10.0 * sum(signal in normalized_document for signal in signals)
            scored.append((index, score))
        return sorted(scored, key=lambda row: row[1], reverse=True)[:top_n]


class EvidenceBoundBenchmarkRouter:
    """Emit deterministic answers from expected facts already present in evidence.

    This is intentionally not a model-quality test. It exercises the same
    AnswerService JSON parsing, citation-ID, exact-quote, persistence, routing,
    and strict refusal path at corpus scale without masking provider outages.
    """

    def __init__(self, expected_by_question: dict[str, dict[str, Any]]):
        self.expected_by_question = expected_by_question
        self.tier_calls: Counter[str] = Counter()

    @staticmethod
    def _line_quote(content: str, value: str) -> str:
        for line in content.splitlines():
            if value.casefold() in line.casefold():
                return line.strip()
        return value

    def call(
        self,
        db: Session,
        *,
        tier: str,
        system_prompt: str,
        user_prompt: str,
        pinned_model: str | None = None,
        prompt_version: str | None = None,
    ) -> ChatCallResult:
        del db, system_prompt, prompt_version
        question = user_prompt.split("QUESTION:\n", 1)[1].split("\n\nEVIDENCE:\n", 1)[0]
        expected = self.expected_by_question[question]
        blocks = [match.groupdict() for match in EVIDENCE_BLOCK.finditer(user_prompt)]
        expected_filenames = expected.get("expected_filenames", [expected["expected_filename"]])
        expected_blocks = [block for block in blocks if block["file"] in expected_filenames]
        if not expected_blocks or not set(expected_filenames).issubset(
            {block["file"] for block in expected_blocks}
        ):
            payload = {
                "status": "insufficient_evidence",
                "answer": "没有足够证据。",
                "claims": [],
            }
        else:
            claims = []
            for value in expected["answer_contains"]:
                value = str(value)
                block = next(
                    (item for item in expected_blocks if value.casefold() in item["content"].casefold()),
                    None,
                )
                if block is None:
                    payload = {
                        "status": "insufficient_evidence",
                        "answer": "证据块没有包含预期事实。",
                        "claims": [],
                    }
                    break
                claims.append(
                    {
                        "text": value,
                        "evidence": [
                            {
                                "id": block["id"],
                                "quote": self._line_quote(block["content"], value),
                            }
                        ],
                    }
                )
            else:
                payload = {
                    "status": "answered",
                    "answer": "；".join(str(value) for value in expected["answer_contains"]),
                    "claims": claims,
                }
        self.tier_calls[tier] += 1
        return ChatCallResult(
            model_id=pinned_model or f"offline-{tier}",
            request_id=str(uuid.uuid4()),
            content=json.dumps(payload, ensure_ascii=False),
            usage=Usage(),
            latency_ms=0,
        )

ADJECTIVES_ZH = [
    "琥珀", "青铜", "苍穹", "晨曦", "翡翠", "霜白", "金色", "靛蓝", "象牙", "翠绿",
    "赤红", "银色", "紫晶", "珊瑚", "星辉", "云杉", "曜石", "月影", "远山", "流光",
    "静水", "北辰", "南风", "海蓝", "松石", "雪原", "沙金", "墨玉", "天青", "丹霞",
    "极光", "玄铁",
]
ADJECTIVES_EN = [
    "Amber", "Bronze", "Azure", "Dawn", "Emerald", "Frost", "Golden", "Indigo", "Ivory", "Jade",
    "Crimson", "Silver", "Violet", "Coral", "Stellar", "Spruce", "Obsidian", "Lunar", "Alpine", "Luminous",
    "Stillwater", "Polaris", "Zephyr", "Marine", "Turquoise", "Tundra", "Saffron", "Onyx", "Cerulean", "Rosy",
    "Aurora", "Iron",
]
NOUNS_ZH = [
    "猎鹰", "灯塔", "河流", "城堡", "航线", "罗盘", "花园", "桥梁", "港湾", "山谷",
    "飞舟", "橡树", "信标", "水晶", "脉冲", "天琴", "雨燕", "鲸鱼", "星门", "塔楼",
    "方舟", "云雀", "枫叶", "雪豹", "海豚", "隼鸟", "火炬", "棱镜", "轨道", "风帆",
    "湖泊", "峰顶",
]
NOUNS_EN = [
    "Falcon", "Beacon", "River", "Citadel", "Route", "Compass", "Garden", "Bridge", "Harbor", "Valley",
    "Voyager", "Oak", "Signal", "Crystal", "Pulse", "Lyra", "Swift", "Whale", "Stargate", "Tower",
    "Ark", "Lark", "Maple", "Leopard", "Dolphin", "Kestrel", "Torch", "Prism", "Orbit", "Sail",
    "Lake", "Summit",
]
ROLES = ["业务主管", "项目经理", "财务控制员", "风险专员", "交付负责人", "系统所有者"]
ROLE_EN = ["Business Supervisor", "Project Manager", "Financial Controller", "Risk Officer", "Delivery Lead", "System Owner"]
FORMATS = [
    ("md", 0.25),
    ("txt", 0.20),
    ("html", 0.15),
    ("docx", 0.15),
    ("xlsx", 0.15),
    ("pdf", 0.10),
]


@dataclass(slots=True)
class CorpusDocument:
    source_key: str
    path: str
    logical_filename: str
    project: str
    document_type: str
    source_type: str
    lifecycle_status: str
    version_label: str
    entity_zh: str
    entity_en: str
    record_code: str
    approval_threshold: int
    approval_role_zh: str
    approval_role_en: str
    timeout_ms: int
    retention_days: int
    lifecycle_code: str
    acronym: str
    acronym_expansion: str
    expected_sheet: str | None = None
    expected_cell_range: str | None = None


def _format_for(index: int, total: int) -> str:
    cursor = (index - 1) / total
    cumulative = 0.0
    for extension, ratio in FORMATS:
        cumulative += ratio
        if cursor < cumulative:
            return extension
    return "pdf"


def _entity(index: int) -> tuple[str, str]:
    position = index - 1
    adjective = position // len(NOUNS_ZH)
    noun = position % len(NOUNS_ZH)
    return (
        f"{ADJECTIVES_ZH[adjective]}{NOUNS_ZH[noun]}",
        f"{ADJECTIVES_EN[adjective]} {NOUNS_EN[noun]}",
    )


def _record(index: int, *, status: str = "approved", logical_index: int | None = None) -> dict[str, Any]:
    logical = logical_index or index
    entity_zh, entity_en = _entity(logical)
    role_index = logical % len(ROLES)
    threshold = 5000 + logical * 17
    if status == "draft":
        threshold += 777
        role_index = (role_index + 1) % len(ROLES)
    return {
        "record_code": f"KBR-{logical:04d}",
        "entity_zh": entity_zh,
        "entity_en": entity_en,
        "approval_threshold": threshold,
        "approval_role_zh": ROLES[role_index],
        "approval_role_en": ROLE_EN[role_index],
        "timeout_ms": 100 + logical % 71,
        "retention_days": 30 + logical % 19,
        "lifecycle_code": f"TX-{logical:04d}-READY",
        "acronym": f"BKR{logical:04d}",
        "acronym_expansion": f"Business Knowledge Rule {logical:04d}",
    }


def _source_type(extension: str, logical_index: int) -> str:
    if extension in {"html", "md"} and logical_index % 3 == 0:
        return "confluence_export"
    if extension in {"docx", "pdf"} and logical_index % 4 == 0:
        return "onenote_export"
    return "upload"


def _body(record: dict[str, Any], project: str, version: str, status: str) -> str:
    return f"""# {record['entity_zh']} / {record['entity_en']} 业务控制规范

## 文档控制 / Document Control

项目：{project}
知识记录编号：{record['record_code']}
版本：{version}
生命周期：{status}

## 正式业务结论 / Approved Business Decision

对于知识记录 {record['record_code']}，{record['entity_zh']}（{record['entity_en']}）业务的取消审批阈值是 CNY {record['approval_threshold']}。当申请金额达到或超过该阈值时，必须由{record['approval_role_zh']}（{record['approval_role_en']}）审批。低于阈值的请求由业务系统按标准流程自动处理。

The canonical service timeout for {record['record_code']} is {record['timeout_ms']} ms. Audit evidence must be retained for {record['retention_days']} days. The authoritative lifecycle code is {record['lifecycle_code']}.

## 术语 / Terminology

缩写 {record['acronym']} 的完整含义是 “{record['acronym_expansion']}”. This acronym is unique to {record['entity_en']} and must not be reused by another business record.

## 非权威说明 / Non-authoritative Note

其他项目可能使用相似的审批流程，但其金额、角色、超时和生命周期代码不得用于 {record['record_code']}。排障手册和会议记录不能覆盖本页的正式结论。
"""


def _write_markdown(path: Path, body: str) -> None:
    path.write_text(body, encoding="utf-8")


def _write_text(path: Path, body: str) -> None:
    text = body.replace("# ", "").replace("## ", "")
    path.write_text(text, encoding="utf-8")


def _write_html(path: Path, record: dict[str, Any], project: str, version: str, status: str) -> None:
    path.write_text(
        f"""<!doctype html><html><body>
<h1>{record['entity_zh']} / {record['entity_en']} 业务控制规范</h1>
<h2>文档控制 / Document Control</h2><p>项目：{project}</p><p>知识记录编号：{record['record_code']}</p><p>版本：{version}</p><p>生命周期：{status}</p>
<h2>正式业务结论 / Approved Business Decision</h2>
<p>对于知识记录 {record['record_code']}，{record['entity_zh']}（{record['entity_en']}）业务的取消审批阈值是 CNY {record['approval_threshold']}。当申请金额达到或超过该阈值时，必须由{record['approval_role_zh']}（{record['approval_role_en']}）审批。</p>
<p>The canonical service timeout for {record['record_code']} is {record['timeout_ms']} ms. Audit evidence must be retained for {record['retention_days']} days. The authoritative lifecycle code is {record['lifecycle_code']}.</p>
<h2>术语 / Terminology</h2><p>缩写 {record['acronym']} 的完整含义是 “{record['acronym_expansion']}”.</p>
<h2>非权威说明</h2><p>其他项目的相似规则不得用于本记录。</p>
</body></html>""",
        encoding="utf-8",
    )


def _write_docx(path: Path, record: dict[str, Any], project: str, version: str, status: str) -> None:
    document = WordDocument()
    document.add_heading(f"{record['entity_zh']} / {record['entity_en']} 业务控制规范", level=1)
    document.add_heading("文档控制 / Document Control", level=2)
    for value in [project, record["record_code"], version, status]:
        document.add_paragraph(str(value))
    document.add_heading("正式业务结论 / Approved Business Decision", level=2)
    document.add_paragraph(
        f"对于知识记录 {record['record_code']}，{record['entity_zh']}（{record['entity_en']}）业务的取消审批阈值是 CNY {record['approval_threshold']}。"
        f"当申请金额达到或超过该阈值时，必须由{record['approval_role_zh']}（{record['approval_role_en']}）审批。"
    )
    table = document.add_table(rows=5, cols=2)
    values = [
        ("Service timeout", f"{record['timeout_ms']} ms"),
        ("Audit retention", f"{record['retention_days']} days"),
        ("Lifecycle code", record["lifecycle_code"]),
        ("Acronym", record["acronym"]),
        ("Expansion", record["acronym_expansion"]),
    ]
    for row, (label, value) in zip(table.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
    document.save(path)


def _write_xlsx(path: Path, record: dict[str, Any], project: str, version: str, status: str) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Approved Rules"
    rows = [
        ("Field", "Authoritative Value"),
        ("Project", project),
        ("Knowledge Record", record["record_code"]),
        ("Entity", f"{record['entity_zh']} / {record['entity_en']}"),
        ("Approval Threshold CNY", record["approval_threshold"]),
        ("Required Approver", f"{record['approval_role_zh']} / {record['approval_role_en']}"),
        ("Service Timeout ms", record["timeout_ms"]),
        ("Audit Retention days", record["retention_days"]),
        ("Lifecycle Code", record["lifecycle_code"]),
        ("Acronym", record["acronym"]),
        ("Acronym Expansion", record["acronym_expansion"]),
        ("Version", version),
        ("Lifecycle", status),
    ]
    for row in rows:
        sheet.append(row)
    sheet.column_dimensions["A"].width = 28
    sheet.column_dimensions["B"].width = 54
    workbook.save(path)


def _pdf_font() -> str:
    candidates = [
        Path("/System/Library/Fonts/PingFang.ttc"),
        Path("/System/Library/Fonts/STHeiti Light.ttc"),
        Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            try:
                pdfmetrics.registerFont(TTFont("BenchmarkCJK", str(candidate), subfontIndex=0))
                return "BenchmarkCJK"
            except Exception:
                continue
    return "Helvetica"


def _write_pdf(path: Path, record: dict[str, Any], project: str, version: str, status: str) -> None:
    canvas = Canvas(str(path))
    font = _pdf_font()
    canvas.setFont(font, 11)
    lines = [line for line in _body(record, project, version, status).replace("#", "").splitlines() if line.strip()]
    y = 800
    for line in lines:
        for start in range(0, len(line), 72):
            canvas.drawString(48, y, line[start : start + 72])
            y -= 17
            if y < 48:
                canvas.showPage()
                canvas.setFont(font, 11)
                y = 800
    canvas.save()


WRITERS = {
    "md": lambda path, record, project, version, status: _write_markdown(path, _body(record, project, version, status)),
    "txt": lambda path, record, project, version, status: _write_text(path, _body(record, project, version, status)),
    "html": _write_html,
    "docx": _write_docx,
    "xlsx": _write_xlsx,
    "pdf": _write_pdf,
}


def generate_corpus(output: Path, count: int, questions: int, seed: int, force: bool) -> dict[str, Any]:
    if count > len(ADJECTIVES_ZH) * len(NOUNS_ZH):
        raise ValueError("The deterministic name space supports at most 1024 logical documents")
    if output.exists() and force:
        shutil.rmtree(output)
    output.mkdir(parents=True, exist_ok=True)
    documents_dir = output / "documents"
    documents_dir.mkdir(exist_ok=True)
    rng = random.Random(seed)
    version_pair_count = min(10, max(1, count // 100))
    deprecated_count = min(10, max(1, count // 100))
    base_count = count - version_pair_count * 2 - deprecated_count
    corpus: list[CorpusDocument] = []

    def create(index: int, logical_index: int, status: str, version: str, extension: str) -> None:
        record = _record(index, status=status, logical_index=logical_index)
        project = f"Project-{(logical_index - 1) % 20 + 1:02d}"
        logical_filename = f"{record['record_code']}-{record['entity_en'].replace(' ', '-')}-control.{extension}"
        disk_name = f"{index:04d}-{status}-{logical_filename}"
        target = documents_dir / disk_name
        WRITERS[extension](target, record, project, version, status)
        corpus.append(
            CorpusDocument(
                source_key=f"source-{index:04d}-{status}",
                path=str(target.relative_to(output)),
                logical_filename=logical_filename,
                project=project,
                document_type=("business-rule", "system-design", "requirement", "runbook")[(logical_index - 1) % 4],
                source_type=_source_type(extension, logical_index),
                lifecycle_status=status,
                version_label=version,
                expected_sheet="Approved Rules" if extension == "xlsx" else None,
                expected_cell_range="A1:B13" if extension == "xlsx" else None,
                **record,
            )
        )

    for index in range(1, base_count + 1):
        create(index, index, "approved", "approved-v1", _format_for(index, count))
    next_index = base_count + 1
    for pair in range(version_pair_count):
        logical_index = base_count + pair + 1
        extension = "md" if pair % 2 == 0 else "docx"
        create(next_index, logical_index, "approved", "approved-v1", extension)
        next_index += 1
        create(next_index, logical_index, "draft", "draft-v2", extension)
        next_index += 1
    for offset in range(deprecated_count):
        logical_index = base_count + version_pair_count + offset + 1
        create(next_index, logical_index, "deprecated", "deprecated-v1", _format_for(logical_index, count))
        next_index += 1

    manifest_path = output / "manifest.jsonl"
    manifest_path.write_text(
        "\n".join(json.dumps(asdict(item), ensure_ascii=False) for item in corpus) + "\n",
        encoding="utf-8",
    )

    approved = [item for item in corpus if item.lifecycle_status == "approved"]
    record_counts = Counter(item.record_code for item in corpus)
    version_pairs = [item for item in approved if record_counts[item.record_code] > 1]
    other_approved = [item for item in approved if record_counts[item.record_code] == 1]
    selected_count = min(questions, len(approved))
    chosen = [
        *version_pairs[:selected_count],
        *rng.sample(other_approved, selected_count - min(len(version_pairs), selected_count)),
    ]
    question_rows: list[dict[str, Any]] = []
    templates = ["exact", "semantic_zh", "cross_language", "acronym", "status"]
    for position, item in enumerate(chosen):
        is_version_pair = record_counts[item.record_code] > 1
        kind = "version" if is_version_pair else templates[(position - len(version_pairs)) % len(templates)]
        if kind == "version":
            question = f"{item.record_code} 的正式 approved 版本规定的取消审批阈值和审批角色是什么？不要采用 draft 值。"
            answer_contains = [str(item.approval_threshold), item.approval_role_zh]
        elif kind == "exact":
            question = f"知识记录 {item.record_code} 的取消审批阈值是多少，必须由什么角色审批？"
            answer_contains = [str(item.approval_threshold), item.approval_role_zh]
        elif kind == "semantic_zh":
            question = f"{item.entity_zh}业务的取消金额达到控制线时，谁负责批准？控制线金额是多少？"
            answer_contains = [str(item.approval_threshold), item.approval_role_zh]
        elif kind == "cross_language":
            question = f"For {item.entity_en}, what are the canonical service timeout and audit retention period?"
            answer_contains = [str(item.timeout_ms), str(item.retention_days)]
        elif kind == "acronym":
            question = f"业务术语 {item.acronym} 的完整英文含义是什么？"
            answer_contains = [item.acronym_expansion]
        else:
            question = f"{item.record_code} 的权威 lifecycle code 是什么？"
            answer_contains = [item.lifecycle_code]
        question_rows.append(
            {
                "id": f"q-{position + 1:04d}",
                "kind": kind,
                "question": question,
                "expected_source_key": item.source_key,
                "expected_filename": item.logical_filename,
                "expected_filenames": [item.logical_filename],
                "expected_version_label": item.version_label,
                "expected_status": "answered",
                "answer_contains": answer_contains,
                "project": item.project,
                "expected_sheet": item.expected_sheet,
                "expected_cell_range": item.expected_cell_range,
            }
        )
    comparison_pool = [item for item in other_approved if item not in chosen]
    for offset in range(min(10, len(comparison_pool) // 2)):
        left = comparison_pool[offset * 2]
        right = comparison_pool[offset * 2 + 1]
        question_rows.append(
            {
                "id": f"q-compare-{offset + 1:03d}",
                "kind": "comparison",
                "question": (
                    f"请比较 {left.record_code} 与 {right.record_code} 的 approved 取消审批阈值，"
                    "分别列出两个金额。"
                ),
                "expected_source_key": None,
                "expected_filename": left.logical_filename,
                "expected_filenames": [left.logical_filename, right.logical_filename],
                "expected_version_label": "approved-v1",
                "expected_status": "answered",
                "answer_contains": [str(left.approval_threshold), str(right.approval_threshold)],
                "project": None,
                "expected_sheet": None,
                "expected_cell_range": None,
            }
        )
    deprecated = [item for item in corpus if item.lifecycle_status == "deprecated"]
    missing_count = max(10, questions // 10)
    for offset in range(missing_count):
        deprecated_item = deprecated[offset] if offset < len(deprecated) else None
        question_rows.append(
            {
                "id": f"q-missing-{offset + 1:03d}",
                "kind": "unanswerable",
                "question": (
                    f"知识记录 {deprecated_item.record_code} 的生产灾备负责人是谁？"
                    if deprecated_item
                    else f"业务术语 BKR{count + 500 + offset:04d} 的完整英文含义是什么？"
                ),
                "expected_source_key": None,
                "expected_filename": None,
                "expected_filenames": [],
                "expected_version_label": None,
                "expected_status": "insufficient_evidence",
                "answer_contains": [],
                "project": None,
            }
        )
    questions_path = output / "questions.jsonl"
    questions_path.write_text(
        "\n".join(json.dumps(row, ensure_ascii=False) for row in question_rows) + "\n",
        encoding="utf-8",
    )
    summary = {
        "generated_at": datetime.now(UTC).isoformat(),
        "seed": seed,
        "documents": len(corpus),
        "approved_documents": len(approved),
        "draft_documents": sum(item.lifecycle_status == "draft" for item in corpus),
        "deprecated_documents": sum(item.lifecycle_status == "deprecated" for item in corpus),
        "questions": len(question_rows),
        "formats": dict(Counter(Path(item.path).suffix.lstrip(".") for item in corpus)),
        "source_types": dict(Counter(item.source_type for item in corpus)),
        "manifest": str(manifest_path),
        "question_file": str(questions_path),
    }
    (output / "generation-summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    return summary


def _read_jsonl(path: Path) -> list[dict[str, Any]]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def load_corpus(corpus_dir: Path, *, reset: bool, embedding: bool) -> dict[str, Any]:
    settings = Settings()
    parser = DocumentParser(settings.parser_backend)
    qwen = qwen_client()
    index = search_index()
    Base.metadata.create_all(engine)
    manifest = _read_jsonl(corpus_dir / "manifest.jsonl")
    started = time.perf_counter()
    parse_times: list[float] = []
    chunks_to_index: list[dict[str, Any]] = []
    format_counts: Counter[str] = Counter()
    warning_counts: Counter[str] = Counter()

    with SessionLocal() as db:
        if reset:
            if index.client.indices.exists(index=settings.search_index_name):
                index.client.indices.delete(index=settings.search_index_name)
            for model in KNOWLEDGE_RESET_ORDER:
                db.execute(delete(model))
            db.commit()
        projects: dict[str, Project] = {}
        for position, row in enumerate(manifest, start=1):
            path = corpus_dir / row["path"]
            format_counts[path.suffix.lower()] += 1
            parse_start = time.perf_counter()
            units, warnings = parser.parse(path)
            parse_times.append((time.perf_counter() - parse_start) * 1000)
            for warning in warnings:
                warning_counts[warning.split(" (")[0]] += 1
            text_chunks = chunk_units(
                units,
                target_tokens=settings.chunk_target_tokens,
                max_tokens=settings.chunk_max_tokens,
                overlap_tokens=settings.chunk_overlap_tokens,
            )
            project = projects.get(row["project"])
            if not project:
                project = db.scalar(select(Project).where(Project.name == row["project"]))
                if not project:
                    project = Project(name=row["project"])
                    db.add(project)
                    db.flush()
                projects[row["project"]] = project
            document = db.scalar(
                select(Document).where(
                    Document.project_id == project.id,
                    Document.filename == row["logical_filename"],
                )
            )
            if not document:
                document = Document(
                    project_id=project.id,
                    logical_key=row["logical_filename"],
                    filename=row["logical_filename"],
                    document_type=row["document_type"],
                    source_type=row["source_type"],
                )
                db.add(document)
                db.flush()
            version = DocumentVersion(
                document_id=document.id,
                sha256=f"benchmark-{row['source_key']}",
                storage_path=row["path"],
                lifecycle_status=row["lifecycle_status"],
                technical_status="searchable",
                is_current=row["lifecycle_status"] == "approved",
                version_label=row["version_label"],
                indexed_at=datetime.now(UTC),
                searchable_at=datetime.now(UTC),
                parse_warnings=warnings,
            )
            db.add(version)
            db.flush()
            for text_chunk in text_chunks:
                chunk = Chunk(
                    version_id=version.id,
                    ordinal=text_chunk.ordinal,
                    heading_path=text_chunk.heading_path,
                    page_number=text_chunk.page_number,
                    sheet_name=text_chunk.sheet_name,
                    cell_range=text_chunk.cell_range,
                    content=text_chunk.content,
                    content_hash=text_chunk.content_hash,
                    record_hash=text_chunk.content_hash,
                    token_count=text_chunk.token_count,
                )
                db.add(chunk)
                db.flush()
                chunks_to_index.append(
                    {
                        "chunk_id": chunk.id,
                        "document_id": document.id,
                        "version_id": version.id,
                        "project_id": project.id,
                        "filename": document.filename,
                        "document_status": version.lifecycle_status,
                        "document_type": document.document_type,
                        "visibility": document.visibility,
                        "version_label": version.version_label,
                        "heading_path": chunk.heading_path,
                        "page_number": chunk.page_number,
                        "sheet_name": chunk.sheet_name,
                        "cell_range": chunk.cell_range,
                        "content": chunk.content,
                    }
                )
            if position % 100 == 0:
                db.commit()
                logger.info("Parsed %s/%s documents", position, len(manifest))
        db.commit()

    cache_hits = 0
    if embedding:
        cache_path = corpus_dir / "embedding-cache.jsonl"
        embedding_cache: dict[str, list[float]] = {}
        if cache_path.exists():
            for row in _read_jsonl(cache_path):
                if len(row.get("embedding", [])) == settings.qwen_embedding_dimensions:
                    embedding_cache[row["content_hash"]] = row["embedding"]
        missing: list[tuple[str, dict[str, Any]]] = []
        for item in chunks_to_index:
            content_hash = hashlib.sha256(item["content"].encode("utf-8")).hexdigest()
            cached = embedding_cache.get(content_hash)
            if cached is not None:
                item["embedding"] = cached
                cache_hits += 1
            else:
                missing.append((content_hash, item))
        with cache_path.open("a", encoding="utf-8") as cache_file:
            for start in range(0, len(missing), 10):
                batch = missing[start : start + 10]
                vectors, _ = qwen.embeddings([item["content"] for _, item in batch])
                for (content_hash, item), vector in zip(batch, vectors):
                    item["embedding"] = vector
                    cache_file.write(
                        json.dumps(
                            {"content_hash": content_hash, "embedding": vector},
                            separators=(",", ":"),
                        )
                        + "\n"
                    )
                cache_file.flush()
                if start and start % 500 == 0:
                    logger.info("Embedded %s/%s missing chunks", start, len(missing))
    index.index_chunks(chunks_to_index, target_index=settings.search_index_name)
    index.activate_alias(settings.search_index_name)
    elapsed = time.perf_counter() - started
    report = {
        "documents": len(manifest),
        "chunks": len(chunks_to_index),
        "embedding_enabled": embedding,
        "embedding_cache_hits": cache_hits,
        "elapsed_seconds": round(elapsed, 3),
        "documents_per_second": round(len(manifest) / elapsed, 2),
        "parse_latency_ms": _latency_stats(parse_times),
        "formats": dict(format_counts),
        "warnings": dict(warning_counts),
        "index": settings.search_index_name,
    }
    (corpus_dir / "load-report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def load_via_production_pipeline_offline(corpus_dir: Path, *, reset: bool) -> dict[str, Any]:
    """Run every benchmark document through the real ingestion job service."""
    settings = Settings(
        storage_root=corpus_dir,
        parser_backend="native",
        allow_bm25_only=True,
        qwen_embedding_model="offline-production-pipeline",
    )
    index = SearchIndex(settings)
    service = IngestionService(
        settings,
        DocumentParser(settings.parser_backend),
        OfflineBenchmarkQwen(),  # type: ignore[arg-type]
        index,
    )
    Base.metadata.create_all(engine)
    manifest = _read_jsonl(corpus_dir / "manifest.jsonl")
    if reset:
        if index.client.indices.exists(index=settings.search_index_name):
            index.client.indices.delete(index=settings.search_index_name)
        with SessionLocal() as db:
            for model in KNOWLEDGE_RESET_ORDER:
                db.execute(delete(model))
            db.commit()
    index.ensure_index()
    index.activate_alias(settings.search_index_name)

    with SessionLocal() as db:
        projects: dict[str, Project] = {}
        for row in manifest:
            project = projects.get(row["project"])
            if not project:
                project = db.scalar(select(Project).where(Project.name == row["project"]))
                if not project:
                    project = Project(name=row["project"])
                    db.add(project)
                    db.flush()
                projects[row["project"]] = project
            document = db.scalar(
                select(Document).where(
                    Document.project_id == project.id,
                    Document.filename == row["logical_filename"],
                )
            )
            if not document:
                document = Document(
                    project_id=project.id,
                    logical_key=row["logical_filename"],
                    filename=row["logical_filename"],
                    document_type=row["document_type"],
                    source_type=row["source_type"],
                )
                db.add(document)
                db.flush()
            version = DocumentVersion(
                document_id=document.id,
                sha256=f"production-benchmark-{row['source_key']}",
                storage_path=row["path"],
                lifecycle_status=row["lifecycle_status"],
                version_label=row["version_label"],
            )
            db.add(version)
            db.flush()
            db.add(IngestionJob(document_id=document.id, version_id=version.id))
        db.commit()

    started = time.perf_counter()
    claim_latencies = []
    process_latencies = []
    processed = 0
    while True:
        with SessionLocal() as db:
            claim_started = time.perf_counter()
            job_id = service.claim_next_job(db)
            claim_latencies.append((time.perf_counter() - claim_started) * 1000)
        if not job_id:
            break
        with SessionLocal() as db:
            process_started = time.perf_counter()
            event_id = service.process(db, job_id)
            service.indexer.publish_event(db, event_id)
            process_latencies.append((time.perf_counter() - process_started) * 1000)
        processed += 1
        if processed % 100 == 0:
            logger.info("Production pipeline processed %s/%s jobs", processed, len(manifest))

    with SessionLocal() as db:
        status_counts = dict(
            db.execute(
                select(IngestionJob.status, func.count()).group_by(IngestionJob.status)
            ).all()
        )
        chunk_count = int(db.scalar(select(func.count()).select_from(Chunk)) or 0)
    report = {
        "mode": "production_ingestion_service_offline_embedding",
        "documents": len(manifest),
        "processed_jobs": processed,
        "job_statuses": status_counts,
        "chunks": chunk_count,
        "elapsed_seconds": round(time.perf_counter() - started, 3),
        "claim_latency_ms": _latency_stats(claim_latencies),
        "process_latency_ms": _latency_stats(process_latencies),
    }
    report["acceptance"] = {
        "passed": processed == len(manifest)
        and status_counts == {"succeeded": len(manifest)}
        and chunk_count > 0,
        "checks": {
            "all_jobs_processed": processed == len(manifest),
            "all_jobs_succeeded": status_counts == {"succeeded": len(manifest)},
            "chunks_created": chunk_count > 0,
        },
    }
    (corpus_dir / "production-ingestion-offline.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return report


def evaluate_api_smoke_offline(corpus_dir: Path) -> dict[str, Any]:
    """Exercise public document/job/source/download APIs with every format."""
    from fastapi.testclient import TestClient

    from app.core.config import get_settings
    from app.main import app

    manifest = _read_jsonl(corpus_dir / "manifest.jsonl")
    samples = {}
    for row in manifest:
        samples.setdefault(Path(row["path"]).suffix.lower(), row)
    project_name = f"Benchmark-API-{uuid.uuid4().hex[:10]}"
    checks = {
        "uploads_accepted": True,
        "duplicate_idempotent": False,
        "jobs_succeeded": True,
        "sources_resolvable": True,
        "downloads_match": True,
    }
    uploaded: list[tuple[dict[str, Any], dict[str, Any]]] = []
    with tempfile.TemporaryDirectory(prefix="knowledge-api-smoke-") as temp_dir:
        settings = Settings(
            storage_root=Path(temp_dir),
            parser_backend="native",
            allow_bm25_only=True,
            qwen_embedding_model="offline-api-smoke",
        )
        index = SearchIndex(settings)
        index.ensure_index()
        index.activate_alias(settings.search_index_name)
        service = IngestionService(
            settings,
            DocumentParser(settings.parser_backend),
            OfflineBenchmarkQwen(),  # type: ignore[arg-type]
            index,
        )
        app.dependency_overrides[get_settings] = lambda: settings
        try:
            with TestClient(app) as client:
                for suffix, row in sorted(samples.items()):
                    path = corpus_dir / row["path"]
                    with path.open("rb") as handle:
                        response = client.post(
                            "/api/v1/documents",
                            files={"file": (path.name, handle, "application/octet-stream")},
                            data={
                                "project": project_name,
                                "document_type": row["document_type"],
                                "lifecycle_status": "approved",
                                "version_label": "api-smoke-v1",
                            },
                        )
                    checks["uploads_accepted"] &= response.status_code == 202
                    uploaded.append((row, response.json()))

                first_row, first_upload = uploaded[0]
                first_path = corpus_dir / first_row["path"]
                with first_path.open("rb") as handle:
                    duplicate = client.post(
                        "/api/v1/documents",
                        files={
                            "file": (first_path.name, handle, "application/octet-stream")
                        },
                        data={
                            "project": project_name,
                            "document_type": first_row["document_type"],
                            "lifecycle_status": "approved",
                            "version_label": "api-smoke-v1",
                        },
                    )
                duplicate_body = duplicate.json()
                checks["duplicate_idempotent"] = (
                    duplicate.status_code == 202
                    and duplicate_body.get("duplicate") is True
                    and duplicate_body.get("version_id") == first_upload["version_id"]
                )

                for _, upload in uploaded:
                    with SessionLocal() as db:
                        event_id = service.process(db, upload["job_id"])
                        service.indexer.publish_event(db, event_id)
                    job_response = client.get(
                        f"/api/v1/ingestion-jobs/{upload['job_id']}"
                    )
                    checks["jobs_succeeded"] &= (
                        job_response.status_code == 200
                        and job_response.json().get("status") == "succeeded"
                    )

                with SessionLocal() as db:
                    chunk_rows = db.execute(
                        select(Chunk, DocumentVersion, Document)
                        .join(DocumentVersion, Chunk.version_id == DocumentVersion.id)
                        .join(Document, DocumentVersion.document_id == Document.id)
                        .where(Document.id.in_([item[1]["document_id"] for item in uploaded]))
                    ).all()
                first_chunks = {}
                for chunk, version, document in chunk_rows:
                    first_chunks.setdefault(document.id, (chunk, version, document))
                for row, upload in uploaded:
                    chunk, version, _ = first_chunks[upload["document_id"]]
                    source = client.get(f"/api/v1/sources/{chunk.id}")
                    checks["sources_resolvable"] &= (
                        source.status_code == 200
                        and source.json().get("document_id") == upload["document_id"]
                    )
                    download = client.get(
                        f"/api/v1/documents/{upload['document_id']}/download",
                        params={"version_id": version.id},
                    )
                    original = (corpus_dir / row["path"]).read_bytes()
                    checks["downloads_match"] &= (
                        download.status_code == 200 and download.content == original
                    )
        finally:
            app.dependency_overrides.pop(get_settings, None)

        if index.client.indices.exists(index=settings.search_index_name):
            index.client.indices.delete(index=settings.search_index_name)

    document_ids = [item[1]["document_id"] for item in uploaded]
    version_ids = [item[1]["version_id"] for item in uploaded]
    with SessionLocal() as db:
        db.execute(delete(Chunk).where(Chunk.version_id.in_(version_ids)))
        db.execute(delete(IngestionJob).where(IngestionJob.document_id.in_(document_ids)))
        db.execute(delete(DocumentVersion).where(DocumentVersion.id.in_(version_ids)))
        db.execute(delete(Document).where(Document.id.in_(document_ids)))
        db.execute(delete(Project).where(Project.name == project_name))
        db.commit()
    SearchIndex(Settings(qwen_embedding_model="offline-production-pipeline")).ensure_index()

    report = {
        "mode": "public_api_offline_smoke",
        "formats": sorted(samples),
        "uploaded_documents": len(uploaded),
        "checks": checks,
        "passed": all(checks.values()) and len(uploaded) == len(samples),
    }
    (corpus_dir / "api-smoke-offline.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return report


def index_existing(corpus_dir: Path, *, embedding: bool) -> dict[str, Any]:
    settings = Settings()
    qwen = qwen_client()
    index = search_index()
    started = time.perf_counter()
    with SessionLocal() as db:
        rows = db.execute(
            select(Chunk, DocumentVersion, Document, Project)
            .join(DocumentVersion, Chunk.version_id == DocumentVersion.id)
            .join(Document, DocumentVersion.document_id == Document.id)
            .join(Project, Document.project_id == Project.id)
            .order_by(Chunk.id)
        ).all()
    chunks_to_index = [
        {
            "chunk_id": chunk.id,
            "document_id": document.id,
            "version_id": version.id,
            "project_id": project.id,
            "filename": document.filename,
            "document_status": version.lifecycle_status,
            "document_type": document.document_type,
            "visibility": document.visibility,
            "version_label": version.version_label,
            "heading_path": chunk.heading_path,
            "page_number": chunk.page_number,
            "sheet_name": chunk.sheet_name,
            "cell_range": chunk.cell_range,
            "content": chunk.content,
        }
        for chunk, version, document, project in rows
    ]
    cache_hits = 0
    cache_path = corpus_dir / "embedding-cache.jsonl"
    embedding_cache = {
        row["content_hash"]: row["embedding"]
        for row in _read_jsonl(cache_path)
        if len(row.get("embedding", [])) == settings.qwen_embedding_dimensions
    } if cache_path.exists() else {}
    missing = []
    if embedding:
        for item in chunks_to_index:
            content_hash = hashlib.sha256(item["content"].encode("utf-8")).hexdigest()
            if content_hash in embedding_cache:
                item["embedding"] = embedding_cache[content_hash]
                cache_hits += 1
            else:
                missing.append((content_hash, item))
        with cache_path.open("a", encoding="utf-8") as cache_file:
            for start in range(0, len(missing), 10):
                batch = missing[start : start + 10]
                vectors, _ = qwen.embeddings([item["content"] for _, item in batch])
                for (content_hash, item), vector in zip(batch, vectors):
                    item["embedding"] = vector
                    cache_file.write(json.dumps({"content_hash": content_hash, "embedding": vector}, separators=(",", ":")) + "\n")
                cache_file.flush()
    if index.client.indices.exists(index=settings.search_index_name):
        index.client.indices.delete(index=settings.search_index_name)
    index.index_chunks(chunks_to_index, target_index=settings.search_index_name)
    index.activate_alias(settings.search_index_name)
    report = {
        "chunks": len(chunks_to_index),
        "embedding_enabled": embedding,
        "embedding_cache_hits": cache_hits,
        "embedding_cache_misses": len(missing),
        "elapsed_seconds": round(time.perf_counter() - started, 3),
        "index": settings.search_index_name,
    }
    (corpus_dir / "index-existing-report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def index_offline_hybrid(corpus_dir: Path) -> dict[str, Any]:
    """Build a local feature-hash vector index without Qwen or network access."""
    settings = Settings(qwen_embedding_model="offline-feature-hash-v1")
    index = SearchIndex(settings)
    started = time.perf_counter()
    with SessionLocal() as db:
        rows = db.execute(
            select(Chunk, DocumentVersion, Document, Project)
            .join(DocumentVersion, Chunk.version_id == DocumentVersion.id)
            .join(Document, DocumentVersion.document_id == Document.id)
            .join(Project, Document.project_id == Project.id)
            .order_by(Chunk.id)
        ).all()
    documents = []
    for chunk, version, document, project in rows:
        documents.append(
            {
                "chunk_id": chunk.id,
                "document_id": document.id,
                "version_id": version.id,
                "project_id": project.id,
                "filename": document.filename,
                "document_status": version.lifecycle_status,
                "document_type": document.document_type,
                "visibility": document.visibility,
                "version_label": version.version_label,
                "heading_path": chunk.heading_path,
                "page_number": chunk.page_number,
                "sheet_name": chunk.sheet_name,
                "cell_range": chunk.cell_range,
                "content": chunk.content,
                "embedding": _offline_feature_vector(chunk.content),
            }
        )
    if index.client.indices.exists(index=settings.search_index_name):
        index.client.indices.delete(index=settings.search_index_name)
    index.index_chunks(documents, target_index=settings.search_index_name)
    index.activate_alias(settings.search_index_name)
    report = {
        "mode": "offline_feature_hash",
        "disclaimer": "Validates Elasticsearch dense-vector mechanics; not Qwen embedding quality.",
        "chunks": len(documents),
        "dimensions": settings.qwen_embedding_dimensions,
        "elapsed_seconds": round(time.perf_counter() - started, 3),
        "index": settings.search_index_name,
    }
    (corpus_dir / "index-offline-hybrid-report.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return report


def _latency_stats(values: list[float]) -> dict[str, float]:
    if not values:
        return {"mean": 0, "p50": 0, "p95": 0, "max": 0}
    ordered = sorted(values)
    return {
        "mean": round(statistics.fmean(values), 3),
        "p50": round(statistics.median(values), 3),
        "p95": round(ordered[min(len(ordered) - 1, math.ceil(len(ordered) * 0.95) - 1)], 3),
        "max": round(max(values), 3),
    }


def _retrieval_acceptance(report: dict[str, Any]) -> dict[str, Any]:
    checks = {
        "recall_at_10_gte_90pct": report["recall_at_10"] >= 0.90,
        "unanswerable_abstention_gte_90pct": (
            report["unanswerable_abstention_rate"] is None
            or report["unanswerable_abstention_rate"] >= 0.90
        ),
        "excel_location_gte_90pct": (
            report["excel_location_accuracy"] is None
            or report["excel_location_accuracy"] >= 0.90
        ),
        "local_p95_lte_2000ms": report["latency_ms"]["p95"] <= 2000,
    }
    if report["mode"]["rerank_requested"]:
        checks["rerank_effective"] = report["mode"]["rerank_effective"]
    return {"passed": all(checks.values()), "checks": checks}


def evaluate_retrieval(
    corpus_dir: Path,
    *,
    limit: int | None,
    use_rerank: bool,
    use_vector: bool,
) -> dict[str, Any]:
    settings = Settings(allow_bm25_only=not use_vector)
    index = search_index()
    qwen = qwen_client()
    retriever = Retriever(settings, index, qwen)
    questions = _read_jsonl(corpus_dir / "questions.jsonl")
    if limit:
        answerable = [row for row in questions if row["expected_status"] == "answered"][:limit]
        missing = [row for row in questions if row["expected_status"] != "answered"][: max(5, limit // 10)]
        questions = [*answerable, *missing]

    original_embedding = retriever._query_embedding
    original_rerank = qwen.rerank
    rerank_successes = 0
    rerank_failures = 0
    if not use_vector:
        def disabled_embedding(_: str) -> list[float]:
            raise RuntimeError("disabled for lexical benchmark")
        retriever._query_embedding = disabled_embedding  # type: ignore[method-assign]
    if not use_rerank:
        def disabled_rerank(*_: Any, **__: Any) -> list[tuple[int, float]]:
            raise QwenAPIError("disabled for RRF benchmark", code="benchmark_disabled")
        qwen.rerank = disabled_rerank  # type: ignore[method-assign]
    else:
        def measured_rerank(*args: Any, **kwargs: Any) -> list[tuple[int, float]]:
            nonlocal rerank_successes, rerank_failures
            try:
                result = original_rerank(*args, **kwargs)
                rerank_successes += 1
                return result
            except QwenAPIError:
                rerank_failures += 1
                raise
        qwen.rerank = measured_rerank  # type: ignore[method-assign]

    latencies = []
    outcomes = []
    hits_at_1 = hits_at_5 = hits_at_10 = 0
    reciprocal_ranks = []
    correct_abstentions = 0
    approved_precedence_total = 0
    approved_precedence_pass = 0
    location_total = 0
    location_pass = 0
    for position, row in enumerate(questions, start=1):
        started = time.perf_counter()
        evidence = retriever.search(row["question"], [])
        latency = (time.perf_counter() - started) * 1000
        latencies.append(latency)
        filenames = [item.filename for item in evidence]
        expected = row.get("expected_filename")
        expected_filenames = row.get("expected_filenames", [expected] if expected else [])
        expected_ranks = [
            filenames.index(filename) + 1 if filename in filenames else None
            for filename in expected_filenames
        ]
        rank = max(expected_ranks) if expected_ranks and all(expected_ranks) else None
        if rank:
            hits_at_1 += rank <= 1
            hits_at_5 += rank <= 5
            hits_at_10 += rank <= 10
            reciprocal_ranks.append(1 / rank)
        elif expected:
            reciprocal_ranks.append(0)
        if row["expected_status"] == "insufficient_evidence" and not evidence:
            correct_abstentions += 1
        if row.get("expected_version_label") and any(item.filename == expected for item in evidence):
            approved_precedence_total += 1
            first_expected = next(item for item in evidence if item.filename == expected)
            approved_precedence_pass += first_expected.version_label == row["expected_version_label"]
        if row.get("expected_sheet") and any(item.filename == expected for item in evidence):
            location_total += 1
            first_expected = next(item for item in evidence if item.filename == expected)
            location_pass += (
                first_expected.sheet_name == row["expected_sheet"]
                and first_expected.cell_range == row["expected_cell_range"]
            )
        outcomes.append(
            {
                "id": row["id"],
                "kind": row["kind"],
                "expected_filename": expected,
                "rank": rank,
                "returned": filenames,
                "latency_ms": round(latency, 3),
            }
        )
        if position % 25 == 0:
            logger.info("Evaluated %s/%s questions", position, len(questions))
    retriever._query_embedding = original_embedding  # type: ignore[method-assign]
    qwen.rerank = original_rerank  # type: ignore[method-assign]
    answerable_count = sum(row["expected_status"] == "answered" for row in questions)
    missing_count = len(questions) - answerable_count
    by_kind: dict[str, list[dict]] = defaultdict(list)
    for result in outcomes:
        by_kind[result["kind"]].append(result)
    report = {
        "mode": {
            "vector_requested": use_vector,
            "rerank_requested": use_rerank,
            "rerank_success_calls": rerank_successes,
            "rerank_failed_calls": rerank_failures,
            "rerank_effective": bool(use_rerank and rerank_successes and not rerank_failures),
        },
        "questions": len(questions),
        "answerable_questions": answerable_count,
        "unanswerable_questions": missing_count,
        "recall_at_1": round(hits_at_1 / answerable_count, 4) if answerable_count else 0,
        "recall_at_5": round(hits_at_5 / answerable_count, 4) if answerable_count else 0,
        "recall_at_10": round(hits_at_10 / answerable_count, 4) if answerable_count else 0,
        "mrr": round(statistics.fmean(reciprocal_ranks), 4) if reciprocal_ranks else 0,
        "unanswerable_abstention_rate": round(correct_abstentions / missing_count, 4) if missing_count else None,
        "approved_precedence_rate": round(approved_precedence_pass / approved_precedence_total, 4) if approved_precedence_total else None,
        "excel_location_accuracy": round(location_pass / location_total, 4) if location_total else None,
        "latency_ms": _latency_stats(latencies),
        "by_kind": {
            kind: {
                "count": len(rows),
                "recall_at_5": round(sum(row["rank"] is not None and row["rank"] <= 5 for row in rows) / len(rows), 4),
            }
            for kind, rows in by_kind.items()
        },
        "failures": [row for row in outcomes if row["expected_filename"] and row["rank"] is None],
    }
    report["acceptance"] = _retrieval_acceptance(report)
    suffix = f"retrieval-v{'1' if use_vector else '0'}-r{'1' if use_rerank else '0'}"
    (corpus_dir / f"{suffix}.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def evaluate_offline_hybrid_retrieval(corpus_dir: Path) -> dict[str, Any]:
    settings = Settings(qwen_embedding_model="offline-feature-hash-v1", allow_bm25_only=False)
    index = SearchIndex(settings)
    qwen = OfflineHybridQwen()
    retriever = Retriever(settings, index, qwen)  # type: ignore[arg-type]
    questions = _read_jsonl(corpus_dir / "questions.jsonl")
    latencies = []
    outcomes = []
    correct_abstentions = 0
    for position, row in enumerate(questions, start=1):
        started = time.perf_counter()
        evidence = retriever.search(row["question"], [])
        latencies.append((time.perf_counter() - started) * 1000)
        filenames = [item.filename for item in evidence]
        expected_filenames = row.get("expected_filenames", [])
        expected_ranks = [
            filenames.index(filename) + 1 if filename in filenames else None
            for filename in expected_filenames
        ]
        rank = max(expected_ranks) if expected_ranks and all(expected_ranks) else None
        if row["expected_status"] == "insufficient_evidence" and not evidence:
            correct_abstentions += 1
        outcomes.append({"id": row["id"], "kind": row["kind"], "rank": rank})
        if position % 25 == 0:
            logger.info("Evaluated offline hybrid retrieval %s/%s", position, len(questions))
    answerable = [row for row in outcomes if not row["id"].startswith("q-missing-")]
    missing_count = len(outcomes) - len(answerable)
    report = {
        "mode": "offline_feature_hash_knn_rrf_rerank",
        "disclaimer": "Validates hybrid mechanics; not Qwen embedding/rerank quality.",
        "questions": len(outcomes),
        "recall_at_1": round(sum(row["rank"] == 1 for row in answerable) / len(answerable), 4),
        "recall_at_5": round(
            sum(row["rank"] is not None and row["rank"] <= 5 for row in answerable)
            / len(answerable),
            4,
        ),
        "recall_at_10": round(
            sum(row["rank"] is not None and row["rank"] <= 10 for row in answerable)
            / len(answerable),
            4,
        ),
        "unanswerable_abstention_rate": round(correct_abstentions / missing_count, 4),
        "latency_ms": _latency_stats(latencies),
        "failures": [row for row in answerable if row["rank"] is None],
    }
    report["acceptance"] = {
        "passed": report["recall_at_10"] >= 0.90
        and report["unanswerable_abstention_rate"] >= 0.90,
        "checks": {
            "recall_at_10_gte_90pct": report["recall_at_10"] >= 0.90,
            "unanswerable_abstention_gte_90pct": (
                report["unanswerable_abstention_rate"] >= 0.90
            ),
        },
    }
    (corpus_dir / "retrieval-offline-hybrid.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return report


def evaluate_project_filtering_offline(corpus_dir: Path) -> dict[str, Any]:
    settings = Settings(qwen_embedding_model="offline-production-pipeline", allow_bm25_only=True)
    retriever = Retriever(settings, SearchIndex(settings), OfflineBenchmarkQwen())  # type: ignore[arg-type]
    rows = [
        row
        for row in _read_jsonl(corpus_dir / "questions.jsonl")
        if row["expected_status"] == "answered" and row.get("project")
    ]
    with SessionLocal() as db:
        project_ids = {project.name: project.id for project in db.scalars(select(Project))}
    outcomes = []
    for row in rows:
        evidence = retriever.search(row["question"], [project_ids[row["project"]]])
        expected_filenames = set(row.get("expected_filenames", []))
        filenames = {item.filename for item in evidence}
        outcomes.append(
            {
                "id": row["id"],
                "expected_hit": expected_filenames.issubset(filenames),
                "all_evidence_in_project": all(
                    item.project_id == project_ids[row["project"]] for item in evidence
                ),
            }
        )
    report = {
        "mode": "project_filtering_offline",
        "questions": len(outcomes),
        "hit_accuracy": round(
            sum(row["expected_hit"] for row in outcomes) / len(outcomes), 4
        ),
        "filter_isolation_accuracy": round(
            sum(row["all_evidence_in_project"] for row in outcomes) / len(outcomes), 4
        ),
        "failures": [
            row
            for row in outcomes
            if not row["expected_hit"] or not row["all_evidence_in_project"]
        ],
    }
    report["passed"] = not report["failures"]
    (corpus_dir / "project-filter-offline.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return report


def evaluate_answers(corpus_dir: Path, *, limit: int, model: str) -> dict[str, Any]:
    rows = _read_jsonl(corpus_dir / "questions.jsonl")
    answerable = [row for row in rows if row["expected_status"] == "answered"][:limit]
    missing = [row for row in rows if row["expected_status"] == "insufficient_evidence"][: max(3, limit // 5)]
    selected = [*answerable, *missing]
    service: AnswerService = answer_service()
    results = []
    latencies = []
    status_correct = citation_correct = content_correct = 0
    with SessionLocal() as db:
        for row in selected:
            started = time.perf_counter()
            response = service.answer(
                db,
                question=row["question"],
                project_ids=[],
                conversation_id=None,
                pinned_model=model,
            )
            latency = (time.perf_counter() - started) * 1000
            latencies.append(latency)
            status_ok = response.status == row["expected_status"]
            expected_filenames = set(row.get("expected_filenames", []))
            source_filenames = {source.filename for source in response.sources}
            citations_ok = (
                expected_filenames.issubset(source_filenames)
                if expected_filenames
                else not response.sources and not response.claims
            )
            content_ok = all(value.casefold() in response.answer.casefold() for value in row["answer_contains"])
            status_correct += status_ok
            citation_correct += citations_ok
            content_correct += content_ok
            results.append(
                {
                    "id": row["id"],
                    "question": row["question"],
                    "expected_status": row["expected_status"],
                    "actual_status": response.status,
                    "status_correct": status_ok,
                    "citation_correct": citations_ok,
                    "content_correct": content_ok,
                    "answer": response.answer,
                    "model": response.model_id,
                    "latency_ms": round(latency, 3),
                }
            )
    report = {
        "model": model,
        "questions": len(selected),
        "status_accuracy": round(status_correct / len(selected), 4),
        "citation_accuracy": round(citation_correct / len(selected), 4),
        "content_accuracy": round(content_correct / len(selected), 4),
        "latency_ms": _latency_stats(latencies),
        "results": results,
    }
    report["acceptance"] = {
        "passed": report["status_accuracy"] >= 0.90
        and report["citation_accuracy"] >= 0.95
        and report["content_accuracy"] >= 0.98,
        "checks": {
            "status_accuracy_gte_90pct": report["status_accuracy"] >= 0.90,
            "citation_accuracy_gte_95pct": report["citation_accuracy"] >= 0.95,
            "content_accuracy_gte_98pct": report["content_accuracy"] >= 0.98,
        },
    }
    model_slug = re.sub(r"[^A-Za-z0-9._-]+", "-", model).strip("-") or "unknown-model"
    report_path = corpus_dir / f"answer-report-{model_slug}.json"
    report_path.write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return report


def evaluate_offline_answer_pipeline(corpus_dir: Path) -> dict[str, Any]:
    """Exercise grounded-answer invariants for the complete scale corpus."""
    rows = _read_jsonl(corpus_dir / "questions.jsonl")
    expected_by_question = {row["question"]: row for row in rows}
    settings = Settings(allow_bm25_only=True, validate_citations_against_database=False)
    offline_qwen = OfflineBenchmarkQwen()
    retriever = Retriever(settings, search_index(), offline_qwen)  # type: ignore[arg-type]
    router = EvidenceBoundBenchmarkRouter(expected_by_question)
    service = AnswerService(settings, retriever, router)  # type: ignore[arg-type]
    memory_engine = create_engine("sqlite+pysqlite:///:memory:")
    Base.metadata.create_all(memory_engine)

    outcomes = []
    latencies = []
    status_pass = citation_pass = content_pass = lifecycle_pass = 0
    with Session(memory_engine) as db:
        for position, row in enumerate(rows, start=1):
            started = time.perf_counter()
            response = service.answer(
                db,
                question=row["question"],
                project_ids=[],
                conversation_id=None,
                pinned_model=None,
            )
            latency = (time.perf_counter() - started) * 1000
            latencies.append(latency)
            status_ok = response.status == row["expected_status"]
            expected_filenames = set(row.get("expected_filenames", []))
            if expected_filenames:
                source_filenames = {source.filename for source in response.sources}
                citation_ok = expected_filenames.issubset(source_filenames) and source_filenames.issubset(
                    expected_filenames
                )
                lifecycle_ok = all(
                    source.document_status == "approved" for source in response.sources
                )
            else:
                citation_ok = not response.sources and not response.claims
                lifecycle_ok = not response.sources
            content_ok = all(
                str(value).casefold() in response.answer.casefold()
                for value in row["answer_contains"]
            )
            status_pass += status_ok
            citation_pass += citation_ok
            content_pass += content_ok
            lifecycle_pass += lifecycle_ok
            outcomes.append(
                {
                    "id": row["id"],
                    "kind": row["kind"],
                    "status_ok": status_ok,
                    "citation_ok": citation_ok,
                    "content_ok": content_ok,
                    "lifecycle_ok": lifecycle_ok,
                    "actual_status": response.status,
                    "model": response.model_id,
                    "route_tier": response.route_tier,
                    "source_files": [source.filename for source in response.sources],
                    "latency_ms": round(latency, 3),
                }
            )
            if position % 25 == 0:
                logger.info("Evaluated offline answer pipeline %s/%s", position, len(rows))
    total = len(rows)
    report = {
        "mode": "offline_evidence_bound_pipeline",
        "disclaimer": "Validates answer orchestration and citation invariants, not Qwen quality.",
        "questions": total,
        "status_accuracy": round(status_pass / total, 4),
        "citation_accuracy": round(citation_pass / total, 4),
        "content_accuracy": round(content_pass / total, 4),
        "approved_only_accuracy": round(lifecycle_pass / total, 4),
        "route_tier_calls": dict(router.tier_calls),
        "latency_ms": _latency_stats(latencies),
        "failures": [
            outcome
            for outcome in outcomes
            if not all(
                outcome[key]
                for key in ("status_ok", "citation_ok", "content_ok", "lifecycle_ok")
            )
        ],
    }
    report["acceptance"] = {
        "passed": not report["failures"]
        and report["citation_accuracy"] >= 0.95
        and report["status_accuracy"] >= 0.90,
        "checks": {
            "status_accuracy_gte_90pct": report["status_accuracy"] >= 0.90,
            "citation_accuracy_gte_95pct": report["citation_accuracy"] >= 0.95,
            "approved_only_accuracy_100pct": report["approved_only_accuracy"] == 1.0,
            "no_case_failures": not report["failures"],
        },
    }
    (corpus_dir / "answer-pipeline-offline.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return report


def _print(report: dict[str, Any]) -> None:
    print(json.dumps(report, ensure_ascii=False, indent=2))


def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
    parser = argparse.ArgumentParser(description="Generate and benchmark a deterministic 1K-document RAG corpus")
    sub = parser.add_subparsers(dest="command", required=True)
    generate = sub.add_parser("generate")
    generate.add_argument("--output", type=Path, default=Path("evaluation/generated/benchmark-1k"))
    generate.add_argument("--count", type=int, default=1000)
    generate.add_argument("--questions", type=int, default=200)
    generate.add_argument("--seed", type=int, default=20260812)
    generate.add_argument("--force", action="store_true")
    load = sub.add_parser("load")
    load.add_argument("corpus", type=Path)
    load.add_argument("--reset", action="store_true")
    load.add_argument("--no-embedding", action="store_true")
    index_command = sub.add_parser("index-existing")
    index_command.add_argument("corpus", type=Path)
    index_command.add_argument("--no-embedding", action="store_true")
    retrieve = sub.add_parser("retrieve")
    retrieve.add_argument("corpus", type=Path)
    retrieve.add_argument("--limit", type=int)
    retrieve.add_argument("--no-rerank", action="store_true")
    retrieve.add_argument("--no-vector", action="store_true")
    answers = sub.add_parser("answers")
    answers.add_argument("corpus", type=Path)
    answers.add_argument("--limit", type=int, default=20)
    answers.add_argument("--model", default="qwen3.7-plus-2026-05-26")
    offline_answers = sub.add_parser("answers-offline")
    offline_answers.add_argument("corpus", type=Path)
    offline_hybrid = sub.add_parser("hybrid-offline")
    offline_hybrid.add_argument("corpus", type=Path)
    production_load = sub.add_parser("load-production-offline")
    production_load.add_argument("corpus", type=Path)
    production_load.add_argument("--reset", action="store_true")
    api_smoke = sub.add_parser("api-smoke-offline")
    api_smoke.add_argument("corpus", type=Path)
    project_filter = sub.add_parser("project-filter-offline")
    project_filter.add_argument("corpus", type=Path)
    args = parser.parse_args()
    if args.command == "generate":
        _print(generate_corpus(args.output, args.count, args.questions, args.seed, args.force))
    elif args.command == "load":
        _print(load_corpus(args.corpus, reset=args.reset, embedding=not args.no_embedding))
    elif args.command == "index-existing":
        _print(index_existing(args.corpus, embedding=not args.no_embedding))
    elif args.command == "retrieve":
        report = evaluate_retrieval(
            args.corpus,
            limit=args.limit,
            use_rerank=not args.no_rerank,
            use_vector=not args.no_vector,
        )
        _print(report)
        if not report["acceptance"]["passed"]:
            raise SystemExit(2)
    elif args.command == "answers":
        report = evaluate_answers(args.corpus, limit=args.limit, model=args.model)
        _print(report)
        if not report["acceptance"]["passed"]:
            raise SystemExit(2)
    elif args.command == "answers-offline":
        report = evaluate_offline_answer_pipeline(args.corpus)
        _print(report)
        if not report["acceptance"]["passed"]:
            raise SystemExit(2)
    elif args.command == "hybrid-offline":
        _print(index_offline_hybrid(args.corpus))
        report = evaluate_offline_hybrid_retrieval(args.corpus)
        _print(report)
        if not report["acceptance"]["passed"]:
            raise SystemExit(2)
    elif args.command == "load-production-offline":
        report = load_via_production_pipeline_offline(args.corpus, reset=args.reset)
        _print(report)
        if not report["acceptance"]["passed"]:
            raise SystemExit(2)
    elif args.command == "api-smoke-offline":
        report = evaluate_api_smoke_offline(args.corpus)
        _print(report)
        if not report["passed"]:
            raise SystemExit(2)
    else:
        report = evaluate_project_filtering_offline(args.corpus)
        _print(report)
        if not report["passed"]:
            raise SystemExit(2)


if __name__ == "__main__":
    main()
