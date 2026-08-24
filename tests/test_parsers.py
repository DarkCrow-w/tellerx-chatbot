from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen.canvas import Canvas

from app.knowledge.parsers import DocumentParser


def test_docx_native_preserves_headings_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "design.docx"
    document = Document()
    document.add_heading("订单系统", level=1)
    document.add_paragraph("订单取消需要主管审批。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "状态"
    table.cell(0, 1).text = "含义"
    table.cell(1, 0).text = "APPROVED"
    table.cell(1, 1).text = "已批准"
    document.save(path)

    parser = DocumentParser(backend="native")
    units, warnings = parser.parse(path)
    assert not warnings
    assert any(unit.heading_path == "订单系统" for unit in units)
    assert any(unit.is_table and "APPROVED" in unit.text for unit in units)


def test_excel_parser_preserves_formula_and_range(tmp_path: Path) -> None:
    path = tmp_path / "rules.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "审批规则"
    sheet.append(["类型", "阈值", "合计"])
    sheet.append(["退款", 100, "=B2*2"])
    workbook.save(path)

    parser = DocumentParser(backend="native")
    units, warnings = parser.parse(path)
    assert units[0].sheet_name == "审批规则"
    assert units[0].cell_range == "A1:C2"
    assert "=B2*2" in units[0].text
    assert any("no cached formula value" in warning for warning in warnings)


def test_pdf_parser_rejoins_ascii_identifier_split_by_visual_line_wrap(tmp_path: Path) -> None:
    path = tmp_path / "rule.pdf"
    canvas = Canvas(str(path))
    canvas.drawString(40, 800, "Lifecycle code is TX-0929-R")
    canvas.drawString(40, 780, "EADY.")
    canvas.save()

    units, _ = DocumentParser(backend="native").parse(path)
    assert "TX-0929-READY" in units[0].text


def test_pdf_parser_does_not_join_unrelated_ascii_lines(tmp_path: Path) -> None:
    path = tmp_path / "paragraphs.pdf"
    canvas = Canvas(str(path))
    canvas.drawString(40, 800, "The first requirement")
    canvas.drawString(40, 780, "The second requirement")
    canvas.save()

    units, _ = DocumentParser(backend="native").parse(path)
    assert "requirementThe" not in units[0].text
