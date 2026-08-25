"""Deterministic bilingual corpus and ground-truth generation for benchmarks.

This synthetic fixture is intentionally independent from live customer documents.
Change it only when the benchmark's coverage or question types should change.
"""

from __future__ import annotations

import json
import random
import shutil
from collections import Counter
from dataclasses import asdict, dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document as WordDocument
from openpyxl import Workbook
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen.canvas import Canvas

ADJECTIVES_ZH = [
    "琥珀",
    "青铜",
    "苍穹",
    "晨曦",
    "翡翠",
    "霜白",
    "金色",
    "靛蓝",
    "象牙",
    "翠绿",
    "赤红",
    "银色",
    "紫晶",
    "珊瑚",
    "星辉",
    "云杉",
    "曜石",
    "月影",
    "远山",
    "流光",
    "静水",
    "北辰",
    "南风",
    "海蓝",
    "松石",
    "雪原",
    "沙金",
    "墨玉",
    "天青",
    "丹霞",
    "极光",
    "玄铁",
]
ADJECTIVES_EN = [
    "Amber",
    "Bronze",
    "Azure",
    "Dawn",
    "Emerald",
    "Frost",
    "Golden",
    "Indigo",
    "Ivory",
    "Jade",
    "Crimson",
    "Silver",
    "Violet",
    "Coral",
    "Stellar",
    "Spruce",
    "Obsidian",
    "Lunar",
    "Alpine",
    "Luminous",
    "Stillwater",
    "Polaris",
    "Zephyr",
    "Marine",
    "Turquoise",
    "Tundra",
    "Saffron",
    "Onyx",
    "Cerulean",
    "Rosy",
    "Aurora",
    "Iron",
]
NOUNS_ZH = [
    "猎鹰",
    "灯塔",
    "河流",
    "城堡",
    "航线",
    "罗盘",
    "花园",
    "桥梁",
    "港湾",
    "山谷",
    "飞舟",
    "橡树",
    "信标",
    "水晶",
    "脉冲",
    "天琴",
    "雨燕",
    "鲸鱼",
    "星门",
    "塔楼",
    "方舟",
    "云雀",
    "枫叶",
    "雪豹",
    "海豚",
    "隼鸟",
    "火炬",
    "棱镜",
    "轨道",
    "风帆",
    "湖泊",
    "峰顶",
]
NOUNS_EN = [
    "Falcon",
    "Beacon",
    "River",
    "Citadel",
    "Route",
    "Compass",
    "Garden",
    "Bridge",
    "Harbor",
    "Valley",
    "Voyager",
    "Oak",
    "Signal",
    "Crystal",
    "Pulse",
    "Lyra",
    "Swift",
    "Whale",
    "Stargate",
    "Tower",
    "Ark",
    "Lark",
    "Maple",
    "Leopard",
    "Dolphin",
    "Kestrel",
    "Torch",
    "Prism",
    "Orbit",
    "Sail",
    "Lake",
    "Summit",
]
ROLES = ["业务主管", "项目经理", "财务控制员", "风险专员", "交付负责人", "系统所有者"]
ROLE_EN = [
    "Business Supervisor",
    "Project Manager",
    "Financial Controller",
    "Risk Officer",
    "Delivery Lead",
    "System Owner",
]
FORMATS = [
    ("md", 0.25),
    ("txt", 0.20),
    ("html", 0.15),
    ("docx", 0.15),
    ("xlsx", 0.15),
    ("pdf", 0.10),
]


@dataclass(slots=True)
class CorpusDocument:
    source_key: str
    path: str
    logical_filename: str
    project: str
    document_type: str
    source_type: str
    lifecycle_status: str
    version_label: str
    entity_zh: str
    entity_en: str
    record_code: str
    approval_threshold: int
    approval_role_zh: str
    approval_role_en: str
    timeout_ms: int
    retention_days: int
    lifecycle_code: str
    acronym: str
    acronym_expansion: str
    expected_sheet: str | None = None
    expected_cell_range: str | None = None


def _format_for(index: int, total: int) -> str:
    cursor = (index - 1) / total
    cumulative = 0.0
    for extension, ratio in FORMATS:
        cumulative += ratio
        if cursor < cumulative:
            return extension
    return "pdf"


def _entity(index: int) -> tuple[str, str]:
    position = index - 1
    adjective = position // len(NOUNS_ZH)
    noun = position % len(NOUNS_ZH)
    return (
        f"{ADJECTIVES_ZH[adjective]}{NOUNS_ZH[noun]}",
        f"{ADJECTIVES_EN[adjective]} {NOUNS_EN[noun]}",
    )


def _record(
    index: int, *, status: str = "approved", logical_index: int | None = None
) -> dict[str, Any]:
    logical = logical_index or index
    entity_zh, entity_en = _entity(logical)
    role_index = logical % len(ROLES)
    threshold = 5000 + logical * 17
    if status == "draft":
        threshold += 777
        role_index = (role_index + 1) % len(ROLES)
    return {
        "record_code": f"KBR-{logical:04d}",
        "entity_zh": entity_zh,
        "entity_en": entity_en,
        "approval_threshold": threshold,
        "approval_role_zh": ROLES[role_index],
        "approval_role_en": ROLE_EN[role_index],
        "timeout_ms": 100 + logical % 71,
        "retention_days": 30 + logical % 19,
        "lifecycle_code": f"TX-{logical:04d}-READY",
        "acronym": f"BKR{logical:04d}",
        "acronym_expansion": f"Business Knowledge Rule {logical:04d}",
    }


def _source_type(extension: str, logical_index: int) -> str:
    if extension in {"html", "md"} and logical_index % 3 == 0:
        return "confluence_export"
    if extension in {"docx", "pdf"} and logical_index % 4 == 0:
        return "onenote_export"
    return "upload"


def _body(record: dict[str, Any], project: str, version: str, status: str) -> str:
    return f"""# {record["entity_zh"]} / {record["entity_en"]} 业务控制规范

## 文档控制 / Document Control

项目：{project}
知识记录编号：{record["record_code"]}
版本：{version}
生命周期：{status}

## 正式业务结论 / Approved Business Decision

对于知识记录 {record["record_code"]}，{record["entity_zh"]}（{record["entity_en"]}）业务的取消审批阈值是 CNY {record["approval_threshold"]}。当申请金额达到或超过该阈值时，必须由{record["approval_role_zh"]}（{record["approval_role_en"]}）审批。低于阈值的请求由业务系统按标准流程自动处理。

The canonical service timeout for {record["record_code"]} is {record["timeout_ms"]} ms. Audit evidence must be retained for {record["retention_days"]} days. The authoritative lifecycle code is {record["lifecycle_code"]}.

## 术语 / Terminology

缩写 {record["acronym"]} 的完整含义是 “{record["acronym_expansion"]}”. This acronym is unique to {record["entity_en"]} and must not be reused by another business record.

## 非权威说明 / Non-authoritative Note

其他项目可能使用相似的审批流程，但其金额、角色、超时和生命周期代码不得用于 {record["record_code"]}。排障手册和会议记录不能覆盖本页的正式结论。
"""


def _write_markdown(path: Path, body: str) -> None:
    path.write_text(body, encoding="utf-8")


def _write_text(path: Path, body: str) -> None:
    text = body.replace("# ", "").replace("## ", "")
    path.write_text(text, encoding="utf-8")


def _write_html(
    path: Path, record: dict[str, Any], project: str, version: str, status: str
) -> None:
    path.write_text(
        f"""<!doctype html><html><body>
<h1>{record["entity_zh"]} / {record["entity_en"]} 业务控制规范</h1>
<h2>文档控制 / Document Control</h2><p>项目：{project}</p><p>知识记录编号：{record["record_code"]}</p><p>版本：{version}</p><p>生命周期：{status}</p>
<h2>正式业务结论 / Approved Business Decision</h2>
<p>对于知识记录 {record["record_code"]}，{record["entity_zh"]}（{record["entity_en"]}）业务的取消审批阈值是 CNY {record["approval_threshold"]}。当申请金额达到或超过该阈值时，必须由{record["approval_role_zh"]}（{record["approval_role_en"]}）审批。</p>
<p>The canonical service timeout for {record["record_code"]} is {record["timeout_ms"]} ms. Audit evidence must be retained for {record["retention_days"]} days. The authoritative lifecycle code is {record["lifecycle_code"]}.</p>
<h2>术语 / Terminology</h2><p>缩写 {record["acronym"]} 的完整含义是 “{record["acronym_expansion"]}”.</p>
<h2>非权威说明</h2><p>其他项目的相似规则不得用于本记录。</p>
</body></html>""",
        encoding="utf-8",
    )


def _write_docx(
    path: Path, record: dict[str, Any], project: str, version: str, status: str
) -> None:
    document = WordDocument()
    document.add_heading(f"{record['entity_zh']} / {record['entity_en']} 业务控制规范", level=1)
    document.add_heading("文档控制 / Document Control", level=2)
    for value in [project, record["record_code"], version, status]:
        document.add_paragraph(str(value))
    document.add_heading("正式业务结论 / Approved Business Decision", level=2)
    document.add_paragraph(
        f"对于知识记录 {record['record_code']}，{record['entity_zh']}（{record['entity_en']}）业务的取消审批阈值是 CNY {record['approval_threshold']}。"
        f"当申请金额达到或超过该阈值时，必须由{record['approval_role_zh']}（{record['approval_role_en']}）审批。"
    )
    table = document.add_table(rows=5, cols=2)
    values = [
        ("Service timeout", f"{record['timeout_ms']} ms"),
        ("Audit retention", f"{record['retention_days']} days"),
        ("Lifecycle code", record["lifecycle_code"]),
        ("Acronym", record["acronym"]),
        ("Expansion", record["acronym_expansion"]),
    ]
    for row, (label, value) in zip(table.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
    document.save(path)


def _write_xlsx(
    path: Path, record: dict[str, Any], project: str, version: str, status: str
) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Approved Rules"
    rows = [
        ("Field", "Authoritative Value"),
        ("Project", project),
        ("Knowledge Record", record["record_code"]),
        ("Entity", f"{record['entity_zh']} / {record['entity_en']}"),
        ("Approval Threshold CNY", record["approval_threshold"]),
        ("Required Approver", f"{record['approval_role_zh']} / {record['approval_role_en']}"),
        ("Service Timeout ms", record["timeout_ms"]),
        ("Audit Retention days", record["retention_days"]),
        ("Lifecycle Code", record["lifecycle_code"]),
        ("Acronym", record["acronym"]),
        ("Acronym Expansion", record["acronym_expansion"]),
        ("Version", version),
        ("Lifecycle", status),
    ]
    for row in rows:
        sheet.append(row)
    sheet.column_dimensions["A"].width = 28
    sheet.column_dimensions["B"].width = 54
    workbook.save(path)


def _pdf_font() -> str:
    candidates = [
        Path("/System/Library/Fonts/PingFang.ttc"),
        Path("/System/Library/Fonts/STHeiti Light.ttc"),
        Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            try:
                pdfmetrics.registerFont(TTFont("BenchmarkCJK", str(candidate), subfontIndex=0))
                return "BenchmarkCJK"
            except Exception:  # noqa: BLE001, S112 - optional platform fonts vary
                continue
    return "Helvetica"


def _write_pdf(path: Path, record: dict[str, Any], project: str, version: str, status: str) -> None:
    canvas = Canvas(str(path))
    font = _pdf_font()
    canvas.setFont(font, 11)
    lines = [
        line
        for line in _body(record, project, version, status).replace("#", "").splitlines()
        if line.strip()
    ]
    y = 800
    for line in lines:
        for start in range(0, len(line), 72):
            canvas.drawString(48, y, line[start : start + 72])
            y -= 17
            if y < 48:
                canvas.showPage()
                canvas.setFont(font, 11)
                y = 800
    canvas.save()


WRITERS = {
    "md": lambda path, record, project, version, status: _write_markdown(
        path, _body(record, project, version, status)
    ),
    "txt": lambda path, record, project, version, status: _write_text(
        path, _body(record, project, version, status)
    ),
    "html": _write_html,
    "docx": _write_docx,
    "xlsx": _write_xlsx,
    "pdf": _write_pdf,
}


def generate_corpus(
    output: Path, count: int, questions: int, seed: int, force: bool
) -> dict[str, Any]:
    if count > len(ADJECTIVES_ZH) * len(NOUNS_ZH):
        raise ValueError("The deterministic name space supports at most 1024 logical documents")
    if output.exists() and force:
        shutil.rmtree(output)
    output.mkdir(parents=True, exist_ok=True)
    documents_dir = output / "documents"
    documents_dir.mkdir(exist_ok=True)
    rng = random.Random(seed)
    version_pair_count = min(10, max(1, count // 100))
    deprecated_count = min(10, max(1, count // 100))
    base_count = count - version_pair_count * 2 - deprecated_count
    corpus: list[CorpusDocument] = []

    def create(index: int, logical_index: int, status: str, version: str, extension: str) -> None:
        record = _record(index, status=status, logical_index=logical_index)
        project = f"Project-{(logical_index - 1) % 20 + 1:02d}"
        logical_filename = (
            f"{record['record_code']}-{record['entity_en'].replace(' ', '-')}-control.{extension}"
        )
        disk_name = f"{index:04d}-{status}-{logical_filename}"
        target = documents_dir / disk_name
        WRITERS[extension](target, record, project, version, status)
        corpus.append(
            CorpusDocument(
                source_key=f"source-{index:04d}-{status}",
                path=str(target.relative_to(output)),
                logical_filename=logical_filename,
                project=project,
                document_type=("business-rule", "system-design", "requirement", "runbook")[
                    (logical_index - 1) % 4
                ],
                source_type=_source_type(extension, logical_index),
                lifecycle_status=status,
                version_label=version,
                expected_sheet="Approved Rules" if extension == "xlsx" else None,
                expected_cell_range="A1:B13" if extension == "xlsx" else None,
                **record,
            )
        )

    for index in range(1, base_count + 1):
        create(index, index, "approved", "approved-v1", _format_for(index, count))
    next_index = base_count + 1
    for pair in range(version_pair_count):
        logical_index = base_count + pair + 1
        extension = "md" if pair % 2 == 0 else "docx"
        create(next_index, logical_index, "approved", "approved-v1", extension)
        next_index += 1
        create(next_index, logical_index, "draft", "draft-v2", extension)
        next_index += 1
    for offset in range(deprecated_count):
        logical_index = base_count + version_pair_count + offset + 1
        create(
            next_index,
            logical_index,
            "deprecated",
            "deprecated-v1",
            _format_for(logical_index, count),
        )
        next_index += 1

    manifest_path = output / "manifest.jsonl"
    manifest_path.write_text(
        "\n".join(json.dumps(asdict(item), ensure_ascii=False) for item in corpus) + "\n",
        encoding="utf-8",
    )

    approved = [item for item in corpus if item.lifecycle_status == "approved"]
    record_counts = Counter(item.record_code for item in corpus)
    version_pairs = [item for item in approved if record_counts[item.record_code] > 1]
    other_approved = [item for item in approved if record_counts[item.record_code] == 1]
    selected_count = min(questions, len(approved))
    chosen = [
        *version_pairs[:selected_count],
        *rng.sample(other_approved, selected_count - min(len(version_pairs), selected_count)),
    ]
    question_rows: list[dict[str, Any]] = []
    templates = ["exact", "semantic_zh", "cross_language", "acronym", "status"]
    for position, item in enumerate(chosen):
        is_version_pair = record_counts[item.record_code] > 1
        kind = (
            "version"
            if is_version_pair
            else templates[(position - len(version_pairs)) % len(templates)]
        )
        if kind == "version":
            question = f"{item.record_code} 的正式 approved 版本规定的取消审批阈值和审批角色是什么？不要采用 draft 值。"
            answer_contains = [str(item.approval_threshold), item.approval_role_zh]
        elif kind == "exact":
            question = f"知识记录 {item.record_code} 的取消审批阈值是多少，必须由什么角色审批？"
            answer_contains = [str(item.approval_threshold), item.approval_role_zh]
        elif kind == "semantic_zh":
            question = f"{item.entity_zh}业务的取消金额达到控制线时，谁负责批准？控制线金额是多少？"
            answer_contains = [str(item.approval_threshold), item.approval_role_zh]
        elif kind == "cross_language":
            question = f"For {item.entity_en}, what are the canonical service timeout and audit retention period?"
            answer_contains = [str(item.timeout_ms), str(item.retention_days)]
        elif kind == "acronym":
            question = f"业务术语 {item.acronym} 的完整英文含义是什么？"
            answer_contains = [item.acronym_expansion]
        else:
            question = f"{item.record_code} 的权威 lifecycle code 是什么？"
            answer_contains = [item.lifecycle_code]
        question_rows.append(
            {
                "id": f"q-{position + 1:04d}",
                "kind": kind,
                "question": question,
                "expected_source_key": item.source_key,
                "expected_filename": item.logical_filename,
                "expected_filenames": [item.logical_filename],
                "expected_version_label": item.version_label,
                "expected_status": "answered",
                "answer_contains": answer_contains,
                "project": item.project,
                "expected_sheet": item.expected_sheet,
                "expected_cell_range": item.expected_cell_range,
            }
        )
    comparison_pool = [item for item in other_approved if item not in chosen]
    for offset in range(min(10, len(comparison_pool) // 2)):
        left = comparison_pool[offset * 2]
        right = comparison_pool[offset * 2 + 1]
        question_rows.append(
            {
                "id": f"q-compare-{offset + 1:03d}",
                "kind": "comparison",
                "question": (
                    f"请比较 {left.record_code} 与 {right.record_code} 的 approved 取消审批阈值，"
                    "分别列出两个金额。"
                ),
                "expected_source_key": None,
                "expected_filename": left.logical_filename,
                "expected_filenames": [left.logical_filename, right.logical_filename],
                "expected_version_label": "approved-v1",
                "expected_status": "answered",
                "answer_contains": [str(left.approval_threshold), str(right.approval_threshold)],
                "project": None,
                "expected_sheet": None,
                "expected_cell_range": None,
            }
        )
    deprecated = [item for item in corpus if item.lifecycle_status == "deprecated"]
    missing_count = max(10, questions // 10)
    for offset in range(missing_count):
        deprecated_item = deprecated[offset] if offset < len(deprecated) else None
        question_rows.append(
            {
                "id": f"q-missing-{offset + 1:03d}",
                "kind": "unanswerable",
                "question": (
                    f"知识记录 {deprecated_item.record_code} 的生产灾备负责人是谁？"
                    if deprecated_item
                    else f"业务术语 BKR{count + 500 + offset:04d} 的完整英文含义是什么？"
                ),
                "expected_source_key": None,
                "expected_filename": None,
                "expected_filenames": [],
                "expected_version_label": None,
                "expected_status": "insufficient_evidence",
                "answer_contains": [],
                "project": None,
            }
        )
    questions_path = output / "questions.jsonl"
    questions_path.write_text(
        "\n".join(json.dumps(row, ensure_ascii=False) for row in question_rows) + "\n",
        encoding="utf-8",
    )
    summary = {
        "generated_at": datetime.now(UTC).isoformat(),
        "seed": seed,
        "documents": len(corpus),
        "approved_documents": len(approved),
        "draft_documents": sum(item.lifecycle_status == "draft" for item in corpus),
        "deprecated_documents": sum(item.lifecycle_status == "deprecated" for item in corpus),
        "questions": len(question_rows),
        "formats": dict(Counter(Path(item.path).suffix.lstrip(".") for item in corpus)),
        "source_types": dict(Counter(item.source_type for item in corpus)),
        "manifest": str(manifest_path),
        "question_file": str(questions_path),
    }
    (output / "generation-summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return summary


def _read_jsonl(path: Path) -> list[dict[str, Any]]:
    return [
        json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()
    ]
